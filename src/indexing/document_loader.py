"""
Document Loader for RAG Pipeline
Supports: PDF, DOCX, TXT, MD, JSON
"""

import json
import logging
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any, Optional

try:
    import markdown
except ImportError:  # pragma: no cover - depends on deployment extras
    markdown = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - depends on deployment extras
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - depends on deployment extras
    DocxDocument = None

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentMetadata:
    """Structure for document metadata"""

    def __init__(
            self,
            filename: str,
            filepath: str,
            file_type: str,
            page_number: Optional[int] = None,
            total_pages: Optional[int] = None,
            timestamp: Optional[str] = None,
            file_size: Optional[int] = None,
            encoding: Optional[str] = None,
            **kwargs
    ):
        self.filename = filename
        self.filepath = filepath
        self.file_type = file_type
        self.page_number = page_number
        self.total_pages = total_pages
        self.timestamp = timestamp or datetime.now().isoformat()
        self.file_size = file_size
        self.encoding = encoding
        self.extra = kwargs  # For any additional metadata

    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary"""
        data = {
            "filename": self.filename,
            "filepath": self.filepath,
            "file_type": self.file_type,
            "timestamp": self.timestamp,
        }
        if self.page_number is not None:
            data["page_number"] = self.page_number
        if self.total_pages is not None:
            data["total_pages"] = self.total_pages
        if self.file_size is not None:
            data["file_size"] = self.file_size
        if self.encoding is not None:
            data["encoding"] = self.encoding
        if self.extra:
            data.update(self.extra)
        return data


class Document:
    """Structure for a loaded document"""

    def __init__(self, content: str, metadata: DocumentMetadata):
        self.content = content
        self.metadata = metadata

    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary"""
        return {
            "content": self.content,
            "metadata": self.metadata.to_dict()
        }


class DocumentLoader:
    """
    Universal document loader supporting multiple file formats.
    Handles PDF, DOCX, TXT, MD, and JSON files.
    """

    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.md', '.json'}

    def __init__(self, default_encoding: str = 'utf-8'):
        """
        Initialize DocumentLoader
        
        Args:
            default_encoding: Default encoding for text files
        """
        self.default_encoding = default_encoding
        logger.info(f"DocumentLoader initialized with encoding: {default_encoding}")

    def load(self, file_path: str) -> List[Document]:
        """
        Load a document from file path
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of Document objects (multiple for multi-page documents)
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        path = Path(file_path)

        # Validate file exists
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file format
        file_ext = path.suffix.lower()
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {self.SUPPORTED_FORMATS}"
            )

        # Get file metadata
        file_size = path.stat().st_size

        logger.info(f"Loading document: {path.name} ({file_ext})")

        try:
            # Route to appropriate loader
            if file_ext == '.pdf':
                documents = self.load_pdf(file_path, file_size)
            elif file_ext == '.docx':
                documents = self.load_docx(file_path, file_size)
            elif file_ext == '.txt':
                documents = self.load_text(file_path, file_size)
            elif file_ext == '.md':
                documents = self.load_markdown(file_path, file_size)
            elif file_ext == '.json':
                documents = self.load_json(file_path, file_size)
            else:
                raise ValueError(f"Unsupported format: {file_ext}")

            logger.info(f"Successfully loaded {len(documents)} document(s) from {path.name}")
            return documents

        except Exception as e:
            logger.error(f"Error loading {file_path}: {str(e)}")
            raise

    @staticmethod
    def load_pdf(file_path: str, file_size: int) -> List[Document]:
        """
        Load PDF document
        
        Args:
            file_path: Path to PDF file
            file_size: Size of file in bytes
            
        Returns:
            List of Document objects (one per page)
        """
        path = Path(file_path)
        documents = []

        try:
            if PdfReader is None:
                raise ValueError("pypdf is not installed. Install with: pip install pypdf")

            reader = PdfReader(file_path)
            total_pages = len(reader.pages)

            for page_num, page in enumerate(reader.pages, start=1):
                # Extract text from page
                text = page.extract_text()

                if text.strip():  # Only add pages with content
                    metadata = DocumentMetadata(
                        filename=path.name,
                        filepath=str(path.absolute()),
                        file_type='pdf',
                        page_number=page_num,
                        total_pages=total_pages,
                        file_size=file_size
                    )

                    documents.append(Document(content=text, metadata=metadata))

            if not documents:
                logger.warning(f"No text content found in PDF: {path.name}")

            return documents

        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to load PDF: {str(e)}")

    @staticmethod
    def load_docx(file_path: str, file_size: int) -> List[Document]:
        """
        Load DOCX document
        
        Args:
            file_path: Path to DOCX file
            file_size: Size of file in bytes
            
        Returns:
            List with single Document object
        """
        path = Path(file_path)

        try:
            if DocxDocument is None:
                raise ValueError("python-docx is not installed. Install with: pip install python-docx")

            doc = DocxDocument(file_path)

            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = '\n\n'.join(paragraphs)

            if not content.strip():
                logger.warning(f"No text content found in DOCX: {path.name}")
                content = ""

            metadata = DocumentMetadata(
                filename=path.name,
                filepath=str(path.absolute()),
                file_type='docx',
                file_size=file_size,
                total_paragraphs=len(paragraphs)
            )

            return [Document(content=content, metadata=metadata)]

        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to load DOCX: {str(e)}")

    def load_text(self, file_path: str, file_size: int) -> List[Document]:
        """
        Load plain text document
        
        Args:
            file_path: Path to text file
            file_size: Size of file in bytes
            
        Returns:
            List with single Document object
        """
        path = Path(file_path)

        # Try multiple encodings
        encodings = [self.default_encoding, 'utf-8', 'latin-1', 'cp1252']
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if content is None:
            raise ValueError(f"Failed to decode text file with tried encodings: {encodings}")

        if not content.strip():
            logger.warning(f"Empty text file: {path.name}")

        metadata = DocumentMetadata(
            filename=path.name,
            filepath=str(path.absolute()),
            file_type='txt',
            file_size=file_size,
            encoding=used_encoding
        )

        return [Document(content=content, metadata=metadata)]

    def load_markdown(self, file_path: str, file_size: int) -> List[Document]:
        """
        Load Markdown document
        
        Args:
            file_path: Path to Markdown file
            file_size: Size of file in bytes
            
        Returns:
            List with single Document object containing both raw and HTML
        """
        path = Path(file_path)

        # Try multiple encodings
        encodings = [self.default_encoding, 'utf-8', 'latin-1', 'cp1252']
        raw_content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    raw_content = f.read()
                used_encoding = encoding
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if raw_content is None:
            raise ValueError(f"Failed to decode markdown file with tried encodings: {encodings}")

        if not raw_content.strip():
            logger.warning(f"Empty markdown file: {path.name}")

        # Convert markdown to HTML (useful for some processing)
        metadata = DocumentMetadata(
            filename=path.name,
            filepath=str(path.absolute()),
            file_type='markdown',
            file_size=file_size,
            encoding=used_encoding,
            has_html=markdown is not None
        )

        # Store raw markdown as content, HTML in metadata
        return [Document(content=raw_content, metadata=metadata)]

    def load_json(self, file_path: str, file_size: int) -> List[Document]:
        """
        Load JSON document
        
        Args:
            file_path: Path to JSON file
            file_size: Size of file in bytes
            
        Returns:
            List of Document objects (handling different JSON structures)
        """
        path = Path(file_path)

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)

            documents = []

            # Handle different JSON structures
            if isinstance(data, list):
                # List of objects
                for idx, item in enumerate(data):
                    content = self._json_to_text(item)
                    metadata = DocumentMetadata(
                        filename=path.name,
                        filepath=str(path.absolute()),
                        file_type='json',
                        file_size=file_size,
                        item_index=idx,
                        total_items=len(data)
                    )
                    documents.append(Document(content=content, metadata=metadata))

            elif isinstance(data, dict):
                # Single object or nested structure
                content = self._json_to_text(data)
                metadata = DocumentMetadata(
                    filename=path.name,
                    filepath=str(path.absolute()),
                    file_type='json',
                    file_size=file_size
                )
                documents.append(Document(content=content, metadata=metadata))

            else:
                # Primitive type (string, number, etc.)
                content = str(data)
                metadata = DocumentMetadata(
                    filename=path.name,
                    filepath=str(path.absolute()),
                    file_type='json',
                    file_size=file_size
                )
                documents.append(Document(content=content, metadata=metadata))

            if not documents:
                logger.warning(f"No content extracted from JSON: {path.name}")

            return documents

        except json.JSONDecodeError as e:
            logger.error(f"Invalid JSON in {file_path}: {str(e)}")
            raise ValueError(f"Failed to parse JSON: {str(e)}")
        except Exception as e:
            logger.error(f"Error reading JSON {file_path}: {str(e)}")
            raise ValueError(f"Failed to load JSON: {str(e)}")

    def _json_to_text(self, obj: Any, indent: int = 0) -> str:
        """
        Convert JSON object to readable text format
        
        Args:
            obj: JSON object (dict, list, or primitive)
            indent: Indentation level
            
        Returns:
            Formatted text string
        """
        lines = []
        prefix = "  " * indent

        if isinstance(obj, dict):
            for key, value in obj.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{prefix}{key}:")
                    lines.append(self._json_to_text(value, indent + 1))
                else:
                    lines.append(f"{prefix}{key}: {value}")

        elif isinstance(obj, list):
            for idx, item in enumerate(obj):
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}[{idx}]:")
                    lines.append(self._json_to_text(item, indent + 1))
                else:
                    lines.append(f"{prefix}- {item}")

        else:
            lines.append(f"{prefix}{obj}")

        return '\n'.join(lines)

    def load_directory(
            self,
            directory_path: str,
            recursive: bool = True,
            file_pattern: Optional[str] = None
    ) -> List[Document]:
        """
        Load all supported documents from a directory
        
        Args:
            directory_path: Path to directory
            recursive: Whether to search subdirectories
            file_pattern: Optional glob pattern (e.g., '*.pdf')
            
        Returns:
            List of all loaded documents
        """
        path = Path(directory_path)

        if not path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        if not path.is_dir():
            raise ValueError(f"Not a directory: {directory_path}")

        all_documents = []

        # Get all files
        if recursive:
            if file_pattern:
                files = path.rglob(file_pattern)
            else:
                files = [f for f in path.rglob('*') if f.is_file()]
        else:
            if file_pattern:
                files = path.glob(file_pattern)
            else:
                files = [f for f in path.glob('*') if f.is_file()]

        # Filter by supported formats
        supported_files = [
            f for f in files
            if f.suffix.lower() in self.SUPPORTED_FORMATS
        ]

        logger.info(f"Found {len(supported_files)} supported files in {directory_path}")

        # Load each file
        for file_path in supported_files:
            try:
                documents = self.load(str(file_path))
                all_documents.extend(documents)
            except Exception as e:
                logger.warning(f"Skipping {file_path.name}: {str(e)}")
                continue

        logger.info(f"Successfully loaded {len(all_documents)} documents from directory")
        return all_documents


# Convenience functions for direct usage
def load_document(file_path: str) -> List[Document]:
    """
    Convenience function to load a single document
    
    Args:
        file_path: Path to document
        
    Returns:
        List of Document objects
    """
    loader = DocumentLoader()
    return loader.load(file_path)


def load_documents_from_directory(
        directory_path: str,
        recursive: bool = True,
        file_pattern: Optional[str] = None
) -> List[Document]:
    """
    Convenience function to load documents from directory
    
    Args:
        directory_path: Path to directory
        recursive: Whether to search subdirectories
        file_pattern: Optional glob pattern
        
    Returns:
        List of all loaded documents
    """
    loader = DocumentLoader()
    return loader.load_directory(directory_path, recursive, file_pattern)
